from pymongo import MongoClient
import tkinter as tk
from tkinter import *
from tkinter import ttk, messagebox
from docx import Document
import datetime
import docx
import os
import sqlite3
import xlsxwriter
from tkinter import filedialog
import smtplib
from email.message import EmailMessage
from validate_email import validate_email

client = MongoClient(
    "mongodb+srv://admin:admin@trackingprogressionsystem-yv211.gcp.mongodb.net/test?retryWrites=true&w=majority")
db = client['Progress_Tracking']
Pro_record = db.Progression_Record
Email_record = db.Email

db_name = 'Recruit_Database.db'


class Progression_Tracking_System(tk.Tk):

    def __init__(self, *args, **kwargs):
        tk.Tk.__init__(self, *args, **kwargs)
        container = tk.Frame(self)

        container.pack(side="top", fill="both", expand=True)
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)

        self.frames = {}

        frame = StartPage(container, self)

        self.frames[StartPage] = frame

        frame.grid(row=0, column=0, sticky="nsew")

        self.show_frame(StartPage)

    def show_frame(self, cont):
        frame = self.frames[cont]
        frame.tkraise()


class StartPage(tk.Frame):

    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)

        self.Date = datetime.datetime.now().date()
        self.Due_Date = datetime.date.today() + datetime.timedelta(days=1)
        self.New_Status = 'New'
        self.Wp_Status = 'WIP'
        self.Delay_Status = 'Delay'

        Label(self, text="File Path: ").grid(row=0, column=0)
        self.File_path = tk.Entry(self, width=40)
        self.File_path.grid(row=0, column=1)

        Label(self, text="Assign To: ").grid(row=0, column=3)
        self.Assign = tk.Entry(self, width=40)
        self.Assign.grid(row=0, column=4)

        Label(self, text="--- Double Click on row to edit the data ---").place(x=10, y=30)

        self.Import_but = tk.Button(self, text="Import", command=self.get_path)
        self.Import_but.grid(row=0, column=2)

        self.Export_but = tk.Button(self, text="Export as Excel", command=self.export_excel)
        self.Export_but.place(x=840, y=20)

        self.refresh_but = tk.Button(self, text="Refresh List", command=self.show_record)
        self.refresh_but.place(x=760, y=20)

        self.Send_email_person = tk.Button(self, text="Send email\nindividually", command=self.SendOne_confirmation)
        self.Send_email_person.place(x=760, y=382)

        self.Send_email_All = tk.Button(self, text="Send email\nAll", command=self.SendAll_confirmation)
        self.Send_email_All.place(x=840, y=382)

        self.Save_but = tk.Button(self, text="Save", command=self.Import_confirmation)
        self.Save_but.grid(row=0, column=5)
        self.set_email_but = tk.Button(self, text="Set Sender's Email", command=self.set_email)
        self.set_email_but.place(x=780, y=430)

        Recoard_frame = LabelFrame(self, text="")
        Recoard_frame.place(x=10, y=50)

        self.tree = ttk.Treeview(Recoard_frame, height=15, column=("1", "2", "3", "4", "5", "6"))
        self.tree.grid(row=1, column=0)
        self.tree.heading('#0', text='Date', anchor=W)
        self.tree.heading(1, text='Role', anchor=W)
        self.tree.heading(2, text='Type', anchor=W)
        self.tree.heading(3, text='Location', anchor=W)
        self.tree.heading(4, text='Assign To', anchor=W)
        self.tree.heading(5, text='Due Date', anchor=W)
        self.tree.heading(6, text='Status', anchor=W)
        self.tree.column('#0', width=100)
        self.tree.column(1, width=200)
        self.tree.column(2, width=110)
        self.tree.column(3, width=100)
        self.tree.column(4, width=200)
        self.tree.column(5, width=100)
        self.tree.column(6, width=100)
        vsb = ttk.Scrollbar(Recoard_frame, orient='vertical', command=self.tree.yview)
        self.tree.grid(row=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        self.tree.configure(yscrollcommand=vsb.set)
        self.tree.bind('<Double-Button>', self.editing)

        Search_frame = LabelFrame(self, text="Searching")
        Search_frame.place(x=10, y=390)

        Label(Search_frame, text="Date:").grid(row=0, column=0)
        self.Date_search = tk.Entry(Search_frame)
        self.Date_search.bind('<KeyRelease>', self.Date_Search)
        self.Date_search.grid(row=0, column=1)

        Label(Search_frame, text="Role:").grid(row=0, column=2)
        self.Role_search = tk.Entry(Search_frame)
        self.Role_search.bind('<KeyRelease>', self.Role_Search)
        self.Role_search.grid(row=0, column=3)

        Label(Search_frame, text="Type:").grid(row=0, column=4)
        self.Type_search = tk.Entry(Search_frame)
        self.Type_search.bind('<KeyRelease>', self.Type_Search)
        self.Type_search.grid(row=0, column=5)

        Label(Search_frame, text="Location:").grid(row=0, column=6)
        self.Location_search = tk.Entry(Search_frame)
        self.Location_search.bind('<KeyRelease>', self.Location_Search)
        self.Location_search.grid(row=0, column=7)

        Label(Search_frame, text="Email:").grid(row=1, column=0)
        self.Email_search = tk.Entry(Search_frame)
        self.Email_search.bind('<KeyRelease>', self.Email_Search)
        self.Email_search.grid(row=1, column=1)

        Label(Search_frame, text="Due_Date:").grid(row=1, column=2)
        self.Due_search = tk.Entry(Search_frame)
        self.Due_search.bind('<KeyRelease>', self.Due_Search)
        self.Due_search.grid(row=1, column=3)

        Label(Search_frame, text="Status:").grid(row=1, column=4)
        self.Status_search = tk.Entry(Search_frame)
        self.Status_search.bind('<KeyRelease>', self.Status_Search)
        self.Status_search.grid(row=1, column=5)

        self.receive_email_address()
        self.update_status()
        self.show_record()

    def receive_email_address(self):
        self.EA = Email_record.find_one({}, {"_id": 0, "Email_Address": 1})
        self.EP = Email_record.find_one({}, {"_id": 0, "Email_Password": 1})
        self.Email_Address = self.EA["Email_Address"]
        self.Email_Password = self.EP["Email_Password"]
        print(self.Email_Address)

    def set_email(self):
        self.sender_main = Toplevel()
        self.sender_main.title('Sender Email')
        self.sender_main.geometry("%dx%d+%d+%d" % (400, 120, 350, 300))

        Label(self.sender_main, text="Current Email: " + self.Email_Address).place(x=70, y=45)

        Label(self.sender_main, text="\tEmail Adrress: ").grid(row=1)
        self.Sender_email = tk.Entry(self.sender_main, width=30)
        self.Sender_email.grid(row=1, column=1)

        Label(self.sender_main, text="\tPassword: ").grid(row=2)
        self.Sender_Pass = tk.Entry(self.sender_main, width=30)
        self.Sender_Pass.grid(row=2, column=1)

        con_but = tk.Button(self.sender_main, text="Save", width=6, height=1, command=self.update_email_address)
        con_but.place(x=140, y=80)
        can_but = tk.Button(self.sender_main, text="Cancel", width=6, height=1, command=self.sender_main.destroy)
        can_but.place(x=200, y=80)

        self.sender_main.focus_set()
        self.sender_main.grab_set()
        self.sender_main.mainloop()

    def update_email_address(self):

        Email_ID = Email_record.find_one({}, {"_id": 1})

        Email_record.update_one(Email_ID, {"$set":
                                               {"Email_Address": self.Sender_email.get(),
                                                "Email_Password": self.Sender_Pass.get()
                                                }
                                           }
                                )
        self.receive_email_address()
        self.sender_main.destroy()

    def SendOne_confirmation(self):
        try:
            self.confirmation = Toplevel()
            self.confirmation.title("Confirm")
            self.confirmation.geometry("%dx%d+%d+%d" % (300, 90, 300, 250))
            recipient = self.tree.item(self.tree.selection())['values'][3]
            Label(self.confirmation, text="Are you sure, you want to send email to\n" + recipient + " ?",
                  font=("Helvetica", 12)).grid(row=0)
            self.confirm_button = Button(self.confirmation, text="Yes", width=5,
                                         command=self.send_email_one)
            self.confirm_button.place(x=80, y=50)
            self.cancel_button = Button(self.confirmation, text="No", width=5,
                                        command=self.confirmation.destroy)
            self.cancel_button.place(x=180, y=50)
            self.confirmation.focus_set()
            self.confirmation.grab_set()
            self.confirmation.mainloop()
        except:
            messagebox.showerror("Error!", "Please select one row")
            self.confirmation.destroy()

    def SendAll_confirmation(self):
        try:
            self.confirmation = Toplevel()
            self.confirmation.title("Confirm")
            self.confirmation.geometry("%dx%d+%d+%d" % (320, 90, 300, 250))
            Label(self.confirmation, text="Are you sure, you want to send email to them?", font=("Helvetica", 12)).grid(
                row=0)
            self.confirm_button = Button(self.confirmation, text="Yes", width=5,
                                         command=self.send_email_all)
            self.confirm_button.place(x=80, y=30)
            self.cancel_button = Button(self.confirmation, text="No", width=5,
                                        command=self.confirmation.destroy)
            self.cancel_button.place(x=180, y=30)
            self.confirmation.focus_set()
            self.confirmation.grab_set()
            self.confirmation.mainloop()
        except:
            messagebox.showerror("Error!", "Something is wrong")
            self.confirmation.destroy()

    def send_email_one(self):
        Date1 = datetime.datetime.strftime(self.Date, '%Y-%m-%d')
        recipient = self.tree.item(self.tree.selection())['values'][3]
        position = self.tree.item(self.tree.selection())['values'][0]
        location = self.tree.item(self.tree.selection())['values'][2]
        due_date = self.tree.item(self.tree.selection())['values'][4]
        with smtplib.SMTP('smtp.office365.com', 587) as smtp:
            smtp.ehlo()
            smtp.starttls()
            smtp.login(self.Email_Address, self.Email_Password)

            subject = 'Reminder on your assignment as of ' + Date1
            line1 = '====================================================================================================================='
            body1 = 'Dear Team,'
            body2 = 'This is automatic email to remind you that the following opportunities are pending to fulfill and they are delayed.\n'
            body3 = '1. ' + position + ', ' + location + ', ' + due_date
            body4 = 'Looking for your cooperation in advance.  If you cannot fulfill within today, please report to your supervisor with the valid reason.'
            body5 = 'Best Regards,'
            body6 = 'Progress Tracking System'
            line2 = '====================================================================================================================='
            msg = f'Subject: {subject}\n\n{line1}\n\n{body1}\n\n{body2}\n\n{body3}\n\n{body4}\n\n{body5}\n\n{body6}\n\n{line2}'

            smtp.sendmail(self.Email_Address, recipient, msg)
        self.confirmation.destroy()

    def send_email_all(self):

        Date1 = datetime.datetime.strftime(self.Date, '%Y-%m-%d')
        content_email = []

        Delay = Pro_record.find({"Status": "Delay"}, {"Assign_To": 1})
        for row in Delay:
            recipient = row["Assign_To"]
            Delay2 = Pro_record.find({"Assign_To": row["Assign_To"], "Status": "Delay"},
                                     {"Role": 1, "Location": 1, "Due_Date": 1})
            for i in Delay2:
                content_email.insert(0, (i["Role"], i["Location"], i["Due_Date"]))

            str1 = '\n'.join(map(str, content_email))
            body3 = str1.replace('(', '').replace(')', '').replace("'", '')
            content_email.clear()
            with smtplib.SMTP('smtp.office365.com', 587) as smtp:

                smtp.ehlo()
                smtp.starttls()
                smtp.login(self.Email_Address, self.Email_Password)
                subject = 'Reminder on your assignment as of ' + Date1
                line1 = '=================================================================='
                body1 = 'Dear Team,'
                body2 = 'This is automatic email to remind you that the following opportunities are pending to fulfill and they are delayed.\n'
                body4 = 'Looking for your cooperation in advance.  If you cannot fulfill within today, please report to your supervisor with the valid reason.'
                body5 = 'Best Regards,'
                body6 = 'Progress Tracking System'
                line2 = '=================================================================='
                msg = f'Subject: {subject}\n\n{line1}\n\n{body1}\n\n{body2}\n\n{body3}\n\n{body4}\n\n{body5}\n\n{body6}\n\n{line2}'

                smtp.sendmail(self.Email_Address, recipient, msg)
        self.confirmation.destroy()

    def Import_confirmation(self):

        self.confirmation = Toplevel()
        self.confirmation.title("Confirm")
        self.confirmation.geometry("%dx%d+%d+%d" % (300, 90, 300, 250))
        Label(self.confirmation, text="Are you sure, you want to import this data?", font=("Helvetica", 12)).grid(row=0)
        self.confirm_button = Button(self.confirmation, text="Yes", width=5,
                                     command=self.Import)
        self.confirm_button.place(x=80, y=30)
        self.cancel_button = Button(self.confirmation, text="No", width=5,
                                    command=self.confirmation.destroy)
        self.cancel_button.place(x=180, y=30)
        self.confirmation.focus_set()
        self.confirmation.grab_set()
        self.confirmation.mainloop()

    def Import(self):
        try:
            file_Path = os.path.abspath(self.File_path.get())
            doc = docx.Document(file_Path)
            Role = doc.paragraphs[0].text
            Para2 = doc.paragraphs[2].text
            new_para2 = re.sub('\s+', '', Para2)
            Type = new_para2.split(":", 1)[1]
            Para3 = doc.paragraphs[3].text
            new_para3 = re.sub('\s+', '', Para3)
            Location = new_para3.split(":", 1)[1]
            if "BTS" in Location:
                BTS = Location[0:3]
                Station = Location[3:]
                Location = BTS + " " + Station

            Email = self.Assign.get()
            Status = "New"
            Record_list = {"Date": str(self.Date),
                           "Role": Role,
                           "Type": Type,
                           "Location": Location,
                           "Assign_To": Email,
                           "Due_Date": str(self.Due_Date),
                           "Status": Status}
            print(Record_list)
            Pro_record.insert_one(Record_list)
            with smtplib.SMTP('smtp.office365.com', 587) as smtp:
                smtp.ehlo()
                smtp.starttls()
                smtp.login(self.Email_Address, self.Email_Password)
                date_str = str(self.Date)
                due_date_str = str(self.Due_Date)
                subject = 'The assignment as of ' + date_str
                line1 = '====================================================================================================================='
                body1 = 'Dear ' + Email
                body2 = 'This is automatic email to assign you an assignment today\n'
                body3 = '1. ' + Role + ', ' + Location + ', ' + due_date_str
                body4 = 'Looking for your cooperation in advance.  If you cannot fulfill within due date time, there will be an email sending to remind you.'
                body5 = 'Best Regards,'
                body6 = 'Progress Tracking System'
                line2 = '====================================================================================================================='
                msg = f'Subject: {subject}\n\n{line1}\n\n{body1}\n\n{body2}\n\n{body3}\n\n{body4}\n\n{body5}\n\n{body6}\n\n{line2}'

                smtp.sendmail(self.Email_Address, Email, msg)

            self.show_record()
            self.File_path.delete(0, 'end')
            self.Assign.delete(0, 'end')
            self.confirmation.destroy()
        except:
            messagebox.showerror("Error!", "Please make sure file path is correct")
            self.confirmation.destroy()

    def editing(self, event):
        try:
            self.tree.item(self.tree.selection())['values'][1]
        except IndexError as e:
            messagebox.showwarning("Error!", "Please select one row")
            return

        old_position = self.tree.item(self.tree.selection())['values'][0]
        old_type = self.tree.item(self.tree.selection())['values'][1]
        old_location = self.tree.item(self.tree.selection())['values'][2]
        old_recipient = self.tree.item(self.tree.selection())['values'][3]
        old_due = self.tree.item(self.tree.selection())['values'][4]
        old_status = self.tree.item(self.tree.selection())['values'][5]

        self.edit_main = Toplevel()
        self.edit_main.title('Edit')
        self.edit_main.geometry('500x500')

        Label(self.edit_main, text='Current Role', font=("Helvetica", 10)).grid(row=0, column=1)
        Pre_Role = Label(self.edit_main, textvariable=StringVar(self.edit_main, value=old_position),
                         font=("Helvetica", 10))
        Pre_Role.grid(row=0, column=2, sticky="w")
        Label(self.edit_main, text='New Role', font=("Helvetica", 14)).grid(row=1, column=1)
        new_role = Text(self.edit_main, height=1, width=30, font=("Helvetica", 14))
        new_role.grid(row=1, column=2)
        new_role['wrap'] = 'none'
        new_role.insert(END, old_position)

        Label(self.edit_main, text='').grid(row=3, columnspan=2)

        Label(self.edit_main, text='Current Type', font=("Helvetica", 10)).grid(row=4, column=1)
        Pre_Type = Label(self.edit_main, textvariable=StringVar(self.edit_main, value=old_type),
                         font=("Helvetica", 10))
        Pre_Type.grid(row=4, column=2, sticky="w")
        Label(self.edit_main, text='New Type', font=("Helvetica", 14)).grid(row=5, column=1)
        new_type = Text(self.edit_main, height=1, width=30, font=("Helvetica", 14))
        new_type.grid(row=5, column=2)
        new_type['wrap'] = 'none'
        new_type.insert(END, old_type)

        Label(self.edit_main, text='').grid(row=6, columnspan=2)

        Label(self.edit_main, text='Current Location', font=("Helvetica", 10)).grid(row=7, column=1)
        Pre_Location = Label(self.edit_main, textvariable=StringVar(self.edit_main, value=old_location),
                             font=("Helvetica", 10))
        Pre_Location.grid(row=7, column=2, sticky="w")
        Label(self.edit_main, text='New Type', font=("Helvetica", 14)).grid(row=8, column=1)
        new_location = Text(self.edit_main, height=1, width=30, font=("Helvetica", 14))
        new_location.grid(row=8, column=2)
        new_location['wrap'] = 'none'
        new_location.insert(END, old_location)

        Label(self.edit_main, text='').grid(row=9, columnspan=2)

        Label(self.edit_main, text='Current Due', font=("Helvetica", 10)).grid(row=10, column=1)
        Pre_Due = Label(self.edit_main, textvariable=StringVar(self.edit_main, value=old_due),
                        font=("Helvetica", 10))
        Pre_Due.grid(row=10, column=2, sticky="w")
        Label(self.edit_main, text='New Due', font=("Helvetica", 14)).grid(row=11, column=1)
        new_due = Text(self.edit_main, height=1, width=30, font=("Helvetica", 14))
        new_due.grid(row=11, column=2)
        new_due['wrap'] = 'none'
        new_due.insert(END, old_due)

        Label(self.edit_main, text='').grid(row=12, columnspan=2)

        Label(self.edit_main, text='Current Email', font=("Helvetica", 10)).grid(row=13, column=1)
        Pre_Email = Label(self.edit_main, textvariable=StringVar(self.edit_main, value=old_recipient),
                          font=("Helvetica", 10))
        Pre_Email.grid(row=13, column=2, sticky="w")
        Label(self.edit_main, text='New Email', font=("Helvetica", 14)).grid(row=14, column=1)
        new_email = Text(self.edit_main, height=1, width=30, font=("Helvetica", 14))
        new_email.grid(row=14, column=2)
        new_email['wrap'] = 'none'
        new_email.insert(END, old_recipient)

        Label(self.edit_main, text='').grid(row=15, columnspan=2)

        Label(self.edit_main, text='Current Status', font=("Helvetica", 10)).grid(row=16, column=1)
        Pre_Status = Label(self.edit_main, textvariable=StringVar(self.edit_main, value=old_status),
                           font=("Helvetica", 10))
        Pre_Status.grid(row=16, column=2, sticky="w")
        Label(self.edit_main, text='New Status', font=("Helvetica", 14)).grid(row=17, column=1)
        new_status = Text(self.edit_main, height=1, width=30, font=("Helvetica", 14))
        new_status.grid(row=17, column=2)
        new_status['wrap'] = 'none'
        new_status.insert(END, old_status)

        Button(self.edit_main, text='Save', font=("Helvetica", 14), width=8, height=1,
               command=lambda: self.edit_record(new_role.get(1.0, 'end-1c'), new_type.get(1.0, 'end-1c'),
                                                new_location.get(1.0, 'end-1c'), new_email.get(1.0, 'end-1c'),
                                                new_due.get(1.0, 'end-1c'), new_status.get(1.0, 'end-1c'),
                                                old_position, old_type, old_location, old_recipient, old_due,
                                                old_status)).place(x=120, y=410)
        Button(self.edit_main, text="Cancel", font=("Helvetica", 14), width=8, height=1,
               command=self.edit_main.destroy).place(x=250, y=410)

        self.edit_main.focus_set()
        self.edit_main.grab_set()
        self.edit_main.mainloop()

    def edit_record(self, new_role, new_type, new_location, new_email, new_due, new_status, old_position, old_type,
                    old_location, old_recipient, old_due, old_status):

        Pro_record.update_one(
            {"Role": old_position, "Type": old_type, "Location": old_location, "Assign_To": old_recipient,
             "Due_Date": old_due, "Status": old_status}, {
                "$set": {"Role": new_role, "Type": new_type, "Location": new_location, "Assign_To": new_email,
                         "Due_Date": new_due, "Status": new_status}})

        self.edit_main.destroy()
        self.show_record()

    def update_status(self):
        Date = datetime.datetime.now().date()
        Delay = "Delay"
        for i in Pro_record.find({}, {"_id": 0, "Due_Date": 1}):
            date_str = datetime.datetime.strptime(i["Due_Date"], '%Y-%m-%d').date()
            if date_str < Date:
                Pro_record.update_many({"Due_Date": i["Due_Date"]}, {"$set": {"Status": Delay}})
        self.show_record()
    def get_path(self):
        word_file = filedialog.askopenfilename()
        self.File_path.insert(0, word_file)

    def export_excel(self):
        excel_path = filedialog.askdirectory()
        list = []
        for i in Pro_record.find({}, {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                      "Status": 1}):
            list.insert(0, (i["Date"], i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))

        workbook = xlsxwriter.Workbook(excel_path + '/Test.xlsx')
        worksheet = workbook.add_worksheet("My sheet")
        top_row = 0
        top_col = 0
        row = 1
        col = 0
        Hcell_format = workbook.add_format()
        Hcell_format.set_pattern()
        Hcell_format.set_bg_color('yellow')
        Hcell_format.set_bold()

        Stcell_format = workbook.add_format()
        Stcell_format.set_pattern()
        Stcell_format.set_bg_color('red')

        for width in range(0, 7):
            worksheet.set_column(width, width, 25)

        head_col = ["Date", "Role", "Type", "Location", "Assign To", "Due Date", "Status"]

        for head in head_col:
            worksheet.write(top_row, top_col, head, Hcell_format)
            top_col += 1

        for date, role, type, location, email, due_date, status in (list):
            worksheet.write(row, col, date)
            worksheet.write(row, col + 1, role)
            worksheet.write(row, col + 2, type)
            worksheet.write(row, col + 3, location)
            worksheet.write(row, col + 4, email)
            worksheet.write(row, col + 5, due_date)
            if status == 'Delay':
                worksheet.write(row, col + 6, status, Stcell_format)
            else:
                worksheet.write(row, col + 6, status)
            row += 1

        workbook.close()

    def show_record(self):
        records = self.tree.get_children()
        for element in records:
            self.tree.delete(element)
        for i in Pro_record.find({}, {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                      "Status": 1}):
            self.tree.insert('', 0, text=i["Date"],
                             values=(i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))

    def Date_Search(self, event):
        records = self.tree.get_children()
        for element in records:
            self.tree.delete(element)
        rgx = re.compile('.*' + self.Date_search.get() + '.*', re.IGNORECASE)
        for i in Pro_record.find({'Date': {'$regex': rgx}},
                                 {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                  "Status": 1}):
            self.tree.insert('', 0, text=i["Date"],
                             values=(i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))

    def Role_Search(self, event):
        records = self.tree.get_children()
        for element in records:
            self.tree.delete(element)
        rgx = re.compile('.*' + self.Role_search.get() + '.*', re.IGNORECASE)
        for i in Pro_record.find({'Role': {'$regex': rgx}},
                                 {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                  "Status": 1}):
            self.tree.insert('', 0, text=i["Date"],
                             values=(i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))

    def Type_Search(self, event):
        records = self.tree.get_children()
        for element in records:
            self.tree.delete(element)
        rgx = re.compile('.*' + self.Type_search.get() + '.*', re.IGNORECASE)
        for i in Pro_record.find({'Type': {'$regex': rgx}},
                                 {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                  "Status": 1}):
            self.tree.insert('', 0, text=i["Date"],
                             values=(i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))

    def Location_Search(self, event):
        records = self.tree.get_children()
        for element in records:
            self.tree.delete(element)
        rgx = re.compile('.*' + self.Location_search.get() + '.*', re.IGNORECASE)
        for i in Pro_record.find({'Location': {'$regex': rgx}},
                                 {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                  "Status": 1}):
            self.tree.insert('', 0, text=i["Date"],
                             values=(i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))

    def Email_Search(self, event):
        records = self.tree.get_children()
        for element in records:
            self.tree.delete(element)
        rgx = re.compile('.*' + self.Email_search.get() + '.*', re.IGNORECASE)
        for i in Pro_record.find({'Assign_To': {'$regex': rgx}},
                                 {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                  "Status": 1}):
            self.tree.insert('', 0, text=i["Date"],
                             values=(i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))

    def Due_Search(self, event):
        records = self.tree.get_children()
        for element in records:
            self.tree.delete(element)
        rgx = re.compile('.*' + self.Due_search.get() + '.*', re.IGNORECASE)
        for i in Pro_record.find({'Due_Date': {'$regex': rgx}},
                                 {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                  "Status": 1}):
            self.tree.insert('', 0, text=i["Date"],
                             values=(i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))

    def Status_Search(self, event):
        records = self.tree.get_children()
        for element in records:
            self.tree.delete(element)
        rgx = re.compile('.*' + self.Status_search.get() + '.*', re.IGNORECASE)
        for i in Pro_record.find({'Status': {'$regex': rgx}},
                                 {"Date": 1, "Role": 1, "Type": 1, "Location": 1, "Assign_To": 1, "Due_Date": 1,
                                  "Status": 1}):
            self.tree.insert('', 0, text=i["Date"],
                             values=(i["Role"], i["Type"], i["Location"], i["Assign_To"], i["Due_Date"], i["Status"]))


app = Progression_Tracking_System()
app.geometry("1000x500")
app.title("Progression Tracking System")
app.mainloop()
