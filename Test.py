import docx
import os


file_Path = os.path.abspath('F:/test/Senior Unix System Admin 20190620.docx')
doc = docx.Document(file_Path)
print(doc.paragraphs[0].text)
print(doc.paragraphs[2].runs[5].text)
Para3 = len(doc.paragraphs[3].runs)
for i in range (2,Para3):
    print(doc.paragraphs[3].runs[i].text)
